"""文件型 BP 材料解析 —— 把上传的 PDF / Word / Excel / 纯文本抽取为可喂入 Deal Intake 分析流的文本。

设计动机（对齐 README「Deal Intake 文件型材料解析」与 agent_design/项目获取Agent.docx）：
自然语言触发（Deal Intake）以消息正文为材料；当用户**上传 BP / 项目表**时，先在 API 层把文件
抽成纯文本，再以同一个 `run_deal_intake(material=..., source_type=...)` 走分析流。本模块只负责
「文件字节 → 文本」这一步，保持纯函数、可独立离线测试、零网络/零数据库。

要点：
- 按扩展名分派（content-type 仅作兜底）；第三方解析库**懒加载**——未安装也不拖垮进程启动/编译，
  仅在真正解析对应格式时抛出可读的依赖缺失错误。
- 体积守卫（默认 20MB）与空文本守卫：抽不出有效文本时给出明确错误，不把空串塞进分析流。
- 不臆造：表格/多页按自然顺序拼接，保留原文，交由下游 LLM 节点做结构化抽取（绝不在这里脑补字段）。
- source_type 推断：Excel→internal_excel（内部项目表），PDF/Word/文本→bp_upload。
"""

from __future__ import annotations

from dataclasses import dataclass, field

from app.objects.deal_list import DealSourceType

# 单文件体积上限（字节）。BP 通常 < 10MB，留一倍冗余；超限早拒绝，避免无谓解析。
MAX_UPLOAD_BYTES = 20 * 1024 * 1024

# 抽取后有效文本下限：低于此视为「没抽到东西」（扫描件图片 PDF / 空表）。
_MIN_TEXT_LEN = 8

# 扩展名 → 规范格式标识
_EXT_FORMAT: dict[str, str] = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".doc": "doc",
    ".xlsx": "xlsx",
    ".xlsm": "xlsx",
    ".xls": "xls",
    ".csv": "csv",
    ".txt": "text",
    ".md": "text",
    ".markdown": "text",
}

# content-type → 规范格式标识（仅当扩展名无法判定时兜底）
_MIME_FORMAT: dict[str, str] = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/msword": "doc",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
    "application/vnd.ms-excel": "xls",
    "text/csv": "csv",
    "text/plain": "text",
    "text/markdown": "text",
}

# 抽取后产物的 source_type 推断（Excel 视作内部项目表，其余视作 BP 上传）
_FORMAT_SOURCE_TYPE: dict[str, DealSourceType] = {
    "xlsx": DealSourceType.INTERNAL_EXCEL,
    "xls": DealSourceType.INTERNAL_EXCEL,
    "csv": DealSourceType.INTERNAL_EXCEL,
}


class DocumentError(Exception):
    """文档解析基类异常（API 层据此返回 4xx，附带 message）。"""


class UnsupportedDocumentError(DocumentError):
    """无法识别/不支持的文件格式。"""


class DocumentTooLargeError(DocumentError):
    """文件超过体积上限。"""


class EmptyDocumentError(DocumentError):
    """抽取后无有效文本（可能是扫描件图片 / 空文件）。"""


class DependencyMissingError(DocumentError):
    """对应格式的解析依赖未安装（部署遗漏，非用户错误）。"""


@dataclass
class ExtractResult:
    text: str
    fmt: str                              # 规范格式标识（pdf/docx/xlsx/...）
    source_type: DealSourceType          # 喂入 run_deal_intake 的来源类型
    unit_count: int = 0                   # 页数（PDF）/ 段落数（Word）/ 工作表数（Excel）
    warnings: list[str] = field(default_factory=list)


def _detect_format(filename: str, content_type: str | None) -> str:
    name = (filename or "").strip().lower()
    dot = name.rfind(".")
    if dot != -1:
        ext = name[dot:]
        if ext in _EXT_FORMAT:
            return _EXT_FORMAT[ext]
    if content_type:
        ct = content_type.split(";", 1)[0].strip().lower()
        if ct in _MIME_FORMAT:
            return _MIME_FORMAT[ct]
    raise UnsupportedDocumentError(
        f"不支持的文件格式：{filename or content_type or '未知'}。"
        "请上传 PDF、Word(.docx)、Excel(.xlsx) 或纯文本(.txt/.md/.csv)。"
    )


def _decode_text(data: bytes) -> str:
    """纯文本解码：UTF-8 优先，失败回退 GB18030（覆盖简繁中文），再不行 latin-1 兜底不抛错。"""
    for enc in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("latin-1", errors="replace")


def _extract_pdf(data: bytes) -> tuple[str, int]:
    try:
        import pdfplumber  # type: ignore
    except ImportError as e:  # pragma: no cover - 依赖缺失路径
        raise DependencyMissingError("PDF 解析依赖 pdfplumber 未安装") from e
    import io

    pages: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            pages.append(page.extract_text() or "")
    return "\n\n".join(p for p in pages if p.strip()), len(pages)


def _extract_docx(data: bytes) -> tuple[str, int]:
    try:
        import docx  # type: ignore  (python-docx)
    except ImportError as e:  # pragma: no cover - 依赖缺失路径
        raise DependencyMissingError("Word 解析依赖 python-docx 未安装") from e
    import io

    document = docx.Document(io.BytesIO(data))
    parts: list[str] = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    # 表格里常含融资条款/团队/财务数据，按行拼接（单元格用制表符分隔，保留原文）
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            line = "\t".join(c for c in cells if c)
            if line.strip():
                parts.append(line)
    return "\n".join(parts), len(parts)


def _extract_xlsx(data: bytes) -> tuple[str, int]:
    try:
        import openpyxl  # type: ignore
    except ImportError as e:  # pragma: no cover - 依赖缺失路径
        raise DependencyMissingError("Excel 解析依赖 openpyxl 未安装") from e
    import io

    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    sheet_count = len(wb.sheetnames)
    sheets: list[str] = []
    for ws in wb.worksheets:
        rows: list[str] = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(v).strip() for v in row if v is not None and str(v).strip()]
            if cells:
                rows.append("\t".join(cells))
        if rows:
            sheets.append(f"# 工作表：{ws.title}\n" + "\n".join(rows))
    wb.close()
    return "\n\n".join(sheets), sheet_count


def _extract_csv(data: bytes) -> tuple[str, int]:
    text = _decode_text(data)
    lines = [ln for ln in text.splitlines() if ln.strip()]
    return "\n".join(lines), len(lines)


def extract_text(
    *, filename: str, data: bytes, content_type: str | None = None
) -> ExtractResult:
    """文件字节 → 可喂入 Deal Intake 的文本。

    抛 DocumentError 子类（体积/格式/空/依赖），API 层据类型映射 4xx/5xx。
    """
    if not data:
        raise EmptyDocumentError("上传文件为空。")
    if len(data) > MAX_UPLOAD_BYTES:
        raise DocumentTooLargeError(
            f"文件超过 {MAX_UPLOAD_BYTES // (1024 * 1024)}MB 上限，请压缩或拆分后再上传。"
        )

    fmt = _detect_format(filename, content_type)
    warnings: list[str] = []

    if fmt in ("doc", "xls"):
        # 旧二进制格式（OLE）解析依赖另一套库，MVP 暂不支持，给出可读引导而非静默失败。
        raise UnsupportedDocumentError(
            "暂不支持旧版 .doc/.xls 二进制格式，请另存为 .docx/.xlsx 后重新上传。"
        )

    if fmt == "text":
        text, unit = _decode_text(data), 0
    elif fmt == "csv":
        text, unit = _extract_csv(data)
    elif fmt == "pdf":
        text, unit = _extract_pdf(data)
        if not text.strip():
            warnings.append("PDF 未抽到文本，可能是扫描件/图片版，建议改传可复制文本的版本。")
    elif fmt == "docx":
        text, unit = _extract_docx(data)
    elif fmt == "xlsx":
        text, unit = _extract_xlsx(data)
    else:  # pragma: no cover - _detect_format 已穷举
        raise UnsupportedDocumentError(f"不支持的格式：{fmt}")

    text = (text or "").strip()
    if len(text) < _MIN_TEXT_LEN:
        raise EmptyDocumentError(
            "未能从文件中抽取到有效文本（可能是扫描件、加密或空文件）。"
            "请确认文件可正常复制文字后重试。"
        )

    source_type = _FORMAT_SOURCE_TYPE.get(fmt, DealSourceType.BP_UPLOAD)
    return ExtractResult(
        text=text, fmt=fmt, source_type=source_type, unit_count=unit, warnings=warnings
    )
